from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


texts: list[str] = [
    "Приложение 5",
    "Сведения об учебно-методическом обеспечении образовательной деятельности юридического лица по заявленным образовательным программам",
    "КЫРГЫЗСКО-ГЕРМАНСКИЙ ИНСТИТУТ ПРИКЛАДНОЙ ИНФОРМАТИКИ",
    "(название юридического лица)",
    "Информатика: Профиль “Веб-информатика”",
    "(название образовательной программы)",
    "Дата заполнения \"____\" ________________ 20__ года",
    "Ректор КГИПИ                          					  А.М. Тологонова",
]

table_texts: list[str] = [
    """    №
    п/п""",
    "Наименование дисциплин учебного плана по курсам обучения",
    "Формы обучения и применяемые технологии",
    "Инвент. ном",
    "Количество учебников       В/ Н  ",
    "Автор",
    "Название",
    "Год издания",
]

table_texts_2 = ["1", "2", "3", "4", "5", "6", "7", "8"]


rows: int = 1
columns: int = len(table_texts)
font_name: str = "Times New Roman"


def set_document_edges(document, size: Inches = Inches(0.5)):
    section = document.sections[0] if document.sections else document.add_section()
    section.left_margin = size
    section.right_margin = size


def add_styled_paragraph(document, text: str, space_before: Inches = None, space_after: Inches = None, is_bold: bool = True, font_size: Pt = Pt(12), font_name: str = "Times New Roman", alignment=WD_ALIGN_PARAGRAPH.CENTER):
    paragraph = document.add_paragraph(text)
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = space_before
    paragraph.paragraph_format.space_after = space_after
    run = paragraph.runs[0]
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = is_bold


def create_table(document, rows: int = rows, columns: int = columns):
    table = document.add_table(rows=rows, cols=columns)
    table.style = "TableGrid"
    table.alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def fill_table_header(table, table_texts: list[str], row_index: int, columns: int = columns):
    for i in range(columns):
        cell = table.rows[row_index].cells[i]
        cell.text = table_texts[i]
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = font_name
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER


def fill_table(row_cells, table_texts: list[str], is_bold: bool = False):
    for i in range(len(table_texts)):
        row_cells[i].text = table_texts[i]
        row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        row_cells[i].paragraphs[0].runs[0].font.bold = is_bold
        row_cells[i].paragraphs[0].runs[0].font.name = font_name
        row_cells[i].paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER


def create_report():
    document = Document()

    set_document_edges(document=document)
    add_styled_paragraph(document=document, text=texts[0], alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_styled_paragraph(document=document, text=texts[1], space_before=Inches(0.2))
    add_styled_paragraph(document=document, text=texts[2], space_after=Inches(0.0))
    add_styled_paragraph(document=document, text=texts[3], font_size=Pt(8), space_before=Inches(0.0))
    add_styled_paragraph(document=document, text=texts[4], space_after=Inches(0.0))
    add_styled_paragraph(document=document, text=texts[5], font_size=Pt(8), space_before=Inches(0.0))

    table = create_table(document=document)
    fill_table_header(table=table, row_index=0, table_texts=table_texts)
    table.add_row()
    fill_table_header(table=table, row_index=1, table_texts=table_texts_2)

    add_styled_paragraph(document=document, text=texts[6], space_before=Inches(1.0), is_bold=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_styled_paragraph(document=document, text=texts[7], space_before=Inches(1.0))

    return document, table

    # document.save("reports/demo.docx")
